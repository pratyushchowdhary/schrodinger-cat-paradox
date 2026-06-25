from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


OUT = "/Users/pcx/Downloads/Check/user-registration-agile-user-stories.docx"


def p(doc, text="", style=None):
    para = doc.add_paragraph(style=style)
    if text:
        para.add_run(text)
    return para


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def story(doc, story_id, title, as_a, i_want, so_that, criteria):
    p(doc, f"{story_id} - {title}", "Heading 2")

    for label, value in [("As a", as_a), ("I want", i_want), ("So that", so_that)]:
        line = p(doc)
        line.paragraph_format.left_indent = Inches(0.15)
        line.paragraph_format.space_after = Pt(2)
        run = line.add_run(f"{label}: ")
        run.bold = True
        line.add_run(value)

    p(doc)
    ac = p(doc)
    ac.add_run("Acceptance criteria").bold = True
    bullets(doc, criteria)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08

for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    s = styles[name]
    s.font.name = "Arial"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(32, 33, 36)
    s.paragraph_format.space_before = Pt(14)
    s.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
run = title.add_run("User Registration - Agile User Stories")
run.font.name = "Arial"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
run.font.size = Pt(26)
run.font.bold = False
run.font.color.rgb = RGBColor(0, 0, 0)

sub = p(doc)
sub.add_run("iHealth and Wellness website | Registration journey").italic = True

p(doc, "Epic", "Heading 1")
p(
    doc,
    "As a new visitor to the iHealth and Wellness website, I want to complete the observed registration and verification flow so that I reach the signed-in profile completion step.",
)

p(doc, "User Journey", "Heading 1")
numbers(
    doc,
    [
        "The user starts on the iHealth and Wellness home page.",
        "The user clicks Sign In.",
        "The user chooses the sign-up path.",
        "The user enters account details.",
        "The registration wizard shows the progress rail: Account Setup, User Type, Security, and Profile.",
        "The user reaches the Who are you joining as? step.",
        "The user selects a user type.",
        "The user continues to the security step.",
        "The user chooses a security question and answer.",
        "The system sends an email verification message.",
        "The user opens the verification email.",
        "The user clicks EMAIL VERIFICATION.",
        "The website opens from the verification link.",
        "The user returns to the website in an authenticated state.",
        "The user sees Sign Out in the header.",
        "The user is prompted to complete their profile.",
    ],
)

p(doc, "Stories", "Heading 1")
stories = [
    (
        "US-REG-01",
        "Start Registration From Sign In",
        "new visitor",
        "to start registration from the website Sign In entry point",
        "I can begin the observed registration flow.",
        [
            "Given I am on the home page, when I click Sign In, then the authentication modal or form is displayed.",
            "Given I choose sign-up, when the registration form opens, then I can begin entering account details.",
        ],
    ),
    (
        "US-REG-02",
        "Enter Account Credentials",
        "new user",
        "to enter my email and password",
        "the registration form can be completed.",
        [
            "Given I am on the registration form, when I enter an email, password, and confirm password, then those fields are populated.",
        ],
    ),
    (
        "US-REG-03",
        "Select User Type",
        "new user",
        "to select a user type",
        "the observed registration wizard can continue.",
        [
            "Given the account setup step is complete, when the wizard shows Who are you joining as?, then the user type options are displayed.",
            "The observed user type options are Patient, Healthcare Professional, Caregiver / Family, Supporter, and Internal User.",
            "The Patient option is shown with the description Living with NF.",
            "The Healthcare Professional option is shown with the description Medical provider or researcher.",
            "The Caregiver / Family option is shown with the description Caring for someone with NF.",
            "The Supporter option is shown with the description NF advocate or donor.",
            "The Internal User option is shown with the description iHealth.",
            "Given a user type is selected, when the user clicks CONTINUE, then the flow proceeds to the next registration step.",
        ],
    ),
    (
        "US-REG-04",
        "Set Security Question",
        "new user",
        "to select a security question and provide an answer",
        "the observed registration form can be completed.",
        [
            "Given I am completing registration, when I open the security question field, then I see a list of available questions.",
            "Given I select a security question, when I enter an answer, then those fields are populated.",
        ],
    ),
    (
        "US-REG-05",
        "Submit Registration",
        "new user",
        "to submit my registration details",
        "I receive the observed verification email.",
        [
            "Given registration details are submitted, then the user receives a verification email.",
        ],
    ),
    (
        "US-REG-06",
        "Receive Verification Email",
        "new user",
        "to receive the verification email",
        "I can continue the observed registration flow.",
        [
            "Given registration is submitted, when I check my inbox, then I receive an email with subject Verify Email Account.",
            "The email is sent from iHealth and Wellness Foundation <noreply@ihealthwellness.org>.",
            "The email contains an EMAIL VERIFICATION call-to-action.",
            "The email explains that verification is required to complete the profile and access resources.",
        ],
    ),
    (
        "US-REG-07",
        "Verify Email",
        "new user",
        "to click the email verification button",
        "the website opens from the verification link.",
        [
            "Given I open the verification email, when I click EMAIL VERIFICATION, then the website opens with a verification token.",
            "Given the website opens from the verification link, then the user returns to the iHealth and Wellness website.",
        ],
    ),
    (
        "US-REG-08",
        "Return To Website After Verification",
        "verified user",
        "to return to the website after verification",
        "I reach the observed signed-in state.",
        [
            "Given I am authenticated, then the header shows Sign Out.",
        ],
    ),
    (
        "US-REG-09",
        "Complete Profile After Registration",
        "newly registered user",
        "to complete my profile",
        "I can continue from the observed profile completion prompt.",
        [
            "Given I am verified and authenticated, when the website loads, then I see a profile completion prompt.",
            "The profile form asks for basic personal and contact information, including first name, last name, email, phone, zip/postal code, and street name.",
            "The profile form includes location-related dropdown behavior, including state/city style selections.",
            "The email field is prefilled with the registered email address.",
        ],
    ),
    (
        "US-REG-10",
        "Access Profile Settings After Registration",
        "newly registered user",
        "to access the site menu after registration",
        "I can see the observed Profile & Settings option.",
        [
            "Given I am signed in, when I open the site menu, then I can see Profile & Settings.",
        ],
    ),
]

for item in stories:
    story(doc, *item)

p(doc, "Evidence From Observed Flow", "Heading 1")
bullets(
    doc,
    [
        "Entry point observed: home page Sign In.",
        "Registration wizard observed: progress rail showed Account Setup, User Type, Security, and Profile; user type screen title was Who are you joining as?; user type options shown were Patient, Healthcare Professional, Caregiver / Family, Supporter, and Internal User; Patient was shown with Living with NF; CONTINUE button was shown on the user type step.",
        "Verification email observed with subject Verify Email Account, sender iHealth and Wellness Foundation <noreply@ihealthwellness.org>, and CTA EMAIL VERIFICATION.",
        "Post-verification state observed: website returned with verification token parameters, header showed Sign Out, and profile completion form appeared.",
        "Profile continuation observed: profile form included personal and contact fields, location-related dropdown behavior was observed, header remained signed in with Sign Out, and opening the menu exposed Profile & Settings.",
    ],
)

p(doc, "Screenshots", "Heading 1")
p(doc, "Email Verification Email", "Heading 2")
p(doc, "Tester-provided screenshot shows the verification email with the EMAIL VERIFICATION CTA. Add the exported screenshot here in Google Docs if needed.")
p(doc, "Post-Verification Profile Prompt", "Heading 2")
p(doc, "Captured screenshot exists but contains visible personal data. Use a redacted version before sharing externally.")

doc.save(OUT)
print(OUT)
