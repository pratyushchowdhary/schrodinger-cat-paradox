from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


OUT = "/Users/pcx/Downloads/Check/user-registration-story-template.docx"


def para(doc, text="", style=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def number(doc, text):
    doc.add_paragraph(text, style="List Number")


def label_line(doc, label, placeholder):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(placeholder)


def story_block(doc, index):
    para(doc, f"US-REG-{index:02d} - [Story Title]", "Heading 2")
    label_line(doc, "As a", "[user/persona]")
    label_line(doc, "I want", "[observed action or goal]")
    label_line(doc, "So that", "[observed outcome only]")
    para(doc)
    p = para(doc)
    p.add_run("Acceptance criteria").bold = True
    bullet(doc, "Given [observed starting state], when [observed user action], then [observed system response].")
    bullet(doc, "Given [observed state], when [observed user action], then [observed result].")
    para(doc)
    p = para(doc)
    p.add_run("Evidence / observation source").bold = True
    bullet(doc, "[Screen, modal, email, menu, URL, or screenshot where this was observed.]")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08

for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    s = styles[name]
    s.font.name = "Arial"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(32, 33, 36)
    s.paragraph_format.space_before = Pt(14)
    s.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
r = title.add_run("User Registration - Story Documentation Template")
r.font.name = "Arial"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0, 0, 0)

subtitle = para(doc)
subtitle.add_run("Use this template to document an observed website registration journey without adding inferred behavior.").italic = True

para(doc, "Instructions", "Heading 1")
bullet(doc, "Document only what was directly observed in the walkthrough.")
bullet(doc, "Do not add validation rules, error states, permissions, token expiry, or alternate paths unless they were observed.")
bullet(doc, "Redact personal data, emails, phone numbers, and tokens.")
bullet(doc, "Use screenshots only when they are redacted and safe to share.")

para(doc, "Feature Context", "Heading 1")
label_line(doc, "Feature", "[Feature name]")
label_line(doc, "Website / product", "[Website or product name]")
label_line(doc, "Observed entry point", "[Example: Home page Sign In]")
label_line(doc, "Observed completion state", "[Example: Signed-in state, profile prompt, confirmation screen]")
label_line(doc, "Recording date", "[Date]")

para(doc, "Epic", "Heading 1")
para(
    doc,
    "As a [user/persona], I want to [complete the observed registration flow] so that [observed end state].",
)

para(doc, "Recorded User Journey", "Heading 1")
for step in [
    "User starts on [starting page].",
    "User clicks [entry point].",
    "User reaches [screen/modal/form].",
    "User enters [observed fields].",
    "User submits [form/action].",
    "System sends or displays [observed confirmation/verification].",
    "User completes [observed next action].",
    "System returns user to [observed page/state].",
    "User sees [observed completion indicator].",
]:
    number(doc, step)

para(doc, "Stories", "Heading 1")
for i in range(1, 4):
    story_block(doc, i)

para(doc, "Observed Evidence", "Heading 1")
bullet(doc, "Entry point observed: [entry point].")
bullet(doc, "Form fields observed: [field list].")
bullet(doc, "Email / verification observed: [subject, sender, CTA, redacted link behavior].")
bullet(doc, "Completion state observed: [state indicator, menu item, modal, success message].")
bullet(doc, "Screenshots captured: [file names or placeholders].")

para(doc, "Screenshots", "Heading 1")
para(doc, "Screenshot 1 - [Screen Name]", "Heading 2")
para(doc, "[Insert redacted screenshot here.]")
para(doc, "Screenshot 2 - [Screen Name]", "Heading 2")
para(doc, "[Insert redacted screenshot here.]")

para(doc, "Out Of Scope / Not Observed", "Heading 1")
bullet(doc, "[List anything intentionally not documented because it was not observed.]")

doc.save(OUT)
print(OUT)
