from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = "/Users/pcx/Downloads/Check/Toni_Use_Case_Specification.docx"


def add_para(doc, text="", style=None):
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    para.add_run(text)
    return para


def add_label(doc, label, value):
    para = doc.add_paragraph()
    run = para.add_run(f"{label}: ")
    run.bold = True
    para.add_run(value)


def add_section_label(doc, label):
    para = doc.add_paragraph()
    run = para.add_run(label)
    run.bold = True
    return para


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_bullets(doc, items):
    if not items:
        doc.add_paragraph("Not specified in PDF")
        return
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_use_case(doc, uc):
    doc.add_paragraph(uc["heading"], style="Heading 2")
    add_label(doc, "Use Case ID", uc["id"])
    add_label(doc, "Use Case Name", uc["name"])
    add_label(doc, "Primary Actor", uc["actor"])
    add_label(doc, "Goal", uc["goal"])
    add_label(doc, "Trigger", uc["trigger"])
    add_label(doc, "Preconditions", uc["preconditions"])
    add_label(doc, "Postconditions", uc["postconditions"])

    add_section_label(doc, "🔄 Main Flow:")
    add_numbered(doc, uc["main"])

    add_section_label(doc, "⚠️ Alternate Flows:")
    add_bullets(doc, uc.get("alternate", []))

    add_section_label(doc, "❌ Exceptions:")
    add_bullets(doc, uc.get("exceptions", []))

    add_section_label(doc, "📋 Business Rules / Notes:")
    add_bullets(doc, uc.get("rules", []))

    add_label(doc, "🔗 Related Screens / Index IDs", ", ".join(uc["nodes"]))


areas = [
    {
        "title": "Functional Area 1 - Authentication and Account Management",
        "use_cases": [
            {
                "id": "UC-01",
                "name": "User Registration",
                "heading": "UC-01 - User Registration",
                "actor": "New User",
                "goal": "Create a user account through the sign-up flow shown in the PDF.",
                "trigger": "User selects Sign Up from the Login / Sign Up page.",
                "preconditions": "User is on the Login / Sign Up page.",
                "postconditions": "Registration form is submitted and verification flow is reached.",
                "main": [
                    "User opens the Login / Sign Up page.",
                    "User reaches the Have an account decision.",
                    "User selects the no/sign-up path.",
                    "System displays the User Registration Form.",
                    "User enters username, email, and password.",
                    "System sends verification email.",
                ],
                "alternate": [
                    "Have an account decision: yes path goes to Login process.",
                    "User may use Google / Facebook sign-up where shown.",
                ],
                "exceptions": [],
                "rules": [
                    "Terms of service, privacy policy, and security policy are shown in the legal step.",
                    "Detailed validation rules are not specified in the PDF.",
                ],
                "nodes": ["01", "02", "03", "04", "05", "06", "08"],
            },
            {
                "id": "UC-02",
                "name": "Email Verification",
                "heading": "UC-02 - Email Verification",
                "actor": "New User",
                "goal": "Complete the email verification flow shown after registration.",
                "trigger": "System sends verification email.",
                "preconditions": "User has submitted registration information.",
                "postconditions": "User completes or fails the verification path shown in the PDF.",
                "main": [
                    "System sends verification email.",
                    "User reaches Email verify decision.",
                    "If verification is yes, system proceeds to account creation/registration confirmation.",
                ],
                "alternate": [
                    "If verification is no, system shows unsuccessful message and instructions.",
                ],
                "exceptions": [],
                "rules": ["The PDF does not specify expiration, retry, or resend rules."],
                "nodes": ["08", "09", "10", "11", "12"],
            },
            {
                "id": "UC-03",
                "name": "Login",
                "heading": "UC-03 - Login",
                "actor": "Registered User",
                "goal": "Log in to the system.",
                "trigger": "User selects Login process.",
                "preconditions": "User has an account.",
                "postconditions": "User reaches the Landing page if information is valid.",
                "main": [
                    "User opens Login process.",
                    "User enters username/email and password.",
                    "User may select Remember me or 2FA option.",
                    "System checks the Information valid decision.",
                    "If information is valid, system logs user into the system.",
                    "System sends user to Landing page.",
                ],
                "alternate": [
                    "User may log in using Google / Facebook.",
                    "If information is not valid, system shows unsuccessful message and instructions.",
                ],
                "exceptions": [],
                "rules": ["2FA is shown as an option only. 2FA setup details are not specified in the PDF."],
                "nodes": ["01", "02", "07", "13", "14", "15", "16", "17", "18", "23"],
            },
            {
                "id": "UC-04",
                "name": "Forgot Password",
                "heading": "UC-04 - Forgot Password",
                "actor": "Registered User",
                "goal": "Use the forgot password form.",
                "trigger": "User selects Forgot password.",
                "preconditions": "User is in the login flow.",
                "postconditions": "Forgot password support flow is presented.",
                "main": [
                    "User selects Forgot password.",
                    "System displays Forgot password form.",
                    "User follows the displayed forgot password flow.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Password reset mechanics are not specified in the PDF."],
                "nodes": ["19", "20", "21", "22"],
            },
            {
                "id": "UC-05",
                "name": "Logout",
                "heading": "UC-05 - Logout",
                "actor": "Registered User",
                "goal": "Log out of the system.",
                "trigger": "User clicks Logout button.",
                "preconditions": "User is logged in.",
                "postconditions": "System shows successfully logout message.",
                "main": [
                    "User clicks logout button.",
                    "System logs the user out.",
                    "System shows successfully logout message.",
                ],
                "alternate": [
                    "System may perform automatic logout because of inactivity timeout.",
                ],
                "exceptions": [],
                "rules": [],
                "nodes": ["119", "120", "121"],
            },
        ],
    },
    {
        "title": "Functional Area 2 - Legal, Profile and Settings",
        "use_cases": [
            {
                "id": "UC-06",
                "name": "Accept Privacy and Health Consent",
                "heading": "UC-06 - Accept Privacy and Health Consent",
                "actor": "Registered User",
                "goal": "Review privacy, ethics, and health data sharing consent items shown in the PDF.",
                "trigger": "User reaches the legal/profile setup flow.",
                "preconditions": "User is registered or logged in.",
                "postconditions": "User proceeds to profile setup.",
                "main": [
                    "System displays Privacy and ethics policy.",
                    "System displays consent for health data sharing.",
                    "User proceeds to User Profile Setup.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Acceptance or rejection mechanics are not specified in the PDF."],
                "nodes": ["24", "25"],
            },
            {
                "id": "UC-07",
                "name": "Create Personal Profile",
                "heading": "UC-07 - Create Personal Profile",
                "actor": "Registered User",
                "goal": "Complete personal information shown in the profile form.",
                "trigger": "User opens Profile section.",
                "preconditions": "User is in User Profile Setup.",
                "postconditions": "Profile information can appear in profile list with information display.",
                "main": [
                    "User opens the personal information form.",
                    "User enters username, photo/avatar, place, date of birth, sex, and health interest.",
                    "User may enter military, education, workplace, language, contact, relationship status, family, details about yourself, name pronunciation, nickname, favorite quotes, blood donations, and life events.",
                    "System shows profile list with information display.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Required vs optional fields are not specified in the PDF."],
                "nodes": ["25", "26", "60"],
            },
            {
                "id": "UC-08",
                "name": "Manage Health Information",
                "heading": "UC-08 - Manage Health Information",
                "actor": "Registered User",
                "goal": "Add or manage health-related profile data shown in the PDF.",
                "trigger": "User enters health data fields in profile setup.",
                "preconditions": "User is in the profile setup/personal information form.",
                "postconditions": "Health information is part of the profile information display.",
                "main": [
                    "User enters health records, health conditions, treatments, symptoms, height, and weight.",
                    "User enters gender identity, ethnicity, race, hospitalizations, insurance, and pronouns.",
                    "System includes the data in the profile list with information display.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Data validation and privacy behavior are not specified in the PDF."],
                "nodes": ["26", "60"],
            },
            {
                "id": "UC-09",
                "name": "Manage Preferences and Visibility",
                "heading": "UC-09 - Manage Preferences and Visibility",
                "actor": "Registered User",
                "goal": "Manage preferences, profile visibility, and language items shown in the PDF.",
                "trigger": "User opens Preferences settings.",
                "preconditions": "User is logged in.",
                "postconditions": "User selects one of the settings items shown.",
                "main": [
                    "User opens Preferences settings.",
                    "User opens Profile visibility.",
                    "User opens Language.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF lists these options but does not specify detailed setting behavior."],
                "nodes": ["27"],
            },
            {
                "id": "UC-10",
                "name": "Manage Account Settings",
                "heading": "UC-10 - Manage Account Settings",
                "actor": "Registered User",
                "goal": "Access account setting options shown in the PDF.",
                "trigger": "User opens Account settings.",
                "preconditions": "User is logged in.",
                "postconditions": "User selects an account setting option.",
                "main": [
                    "User opens Account settings.",
                    "User views username/password option.",
                    "User views linked email option.",
                    "User views delete account option.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Detailed update/delete behavior is not specified in the PDF."],
                "nodes": ["28"],
            },
            {
                "id": "UC-11",
                "name": "Export, Print or Share Data",
                "heading": "UC-11 - Export, Print or Share Data",
                "actor": "Registered User",
                "goal": "Use the export data options shown in the PDF.",
                "trigger": "User opens Export data option.",
                "preconditions": "User has profile or personal data.",
                "postconditions": "User selects print, export to apps, or share data.",
                "main": [
                    "User opens Export data option.",
                    "User selects Print, Export to Apps, or Share data.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Export formats, app destinations, and sharing rules are not specified in the PDF."],
                "nodes": ["29", "1", "2", "3"],
            },
        ],
    },
    {
        "title": "Functional Area 3 - Communities and Events",
        "use_cases": [
            {
                "id": "UC-12",
                "name": "Explore Communities",
                "heading": "UC-12 - Explore Communities",
                "actor": "Registered User",
                "goal": "Explore community options shown in the PDF.",
                "trigger": "User opens Communities section or Explore other communities.",
                "preconditions": "User is logged in.",
                "postconditions": "Community list, search, browse, sort, filter, or recommendations are shown.",
                "main": [
                    "User opens Communities section.",
                    "User opens Explore other communities.",
                    "User may Search by keywords.",
                    "User may Browse by category, friends, groups, or suggestions.",
                    "User may Sort by default, A-Z, recently joined, earliest joined, or popular/most visited.",
                    "User may Filter by place/city, public or not, joined group, or joined status.",
                    "System may Recommend based on joined group, condition, treatment, information, topic, or guidelines.",
                ],
                "alternate": [
                    "System may auto-join community by condition, treatment, or symptom and place user in Joined state.",
                ],
                "exceptions": [],
                "rules": [
                    "Auto-join matching rules are not specified in the PDF.",
                    "The PDF does not specify what happens when search results are empty.",
                ],
                "nodes": ["30", "31", "32", "33", "34", "35", "36", "37", "43"],
            },
            {
                "id": "UC-13",
                "name": "View Community Information",
                "heading": "UC-13 - View Community Information",
                "actor": "Registered User",
                "goal": "View community information displayed in the community list.",
                "trigger": "User selects a community from the list.",
                "preconditions": "Community is visible in list.",
                "postconditions": "Community information is displayed.",
                "main": [
                    "System displays community list with information display.",
                    "System shows group name, topic, number of posts, number of members, and patient stats data.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": [],
                "nodes": ["38", "39"],
            },
            {
                "id": "UC-14",
                "name": "Request Community Access",
                "heading": "UC-14 - Request Community Access",
                "actor": "Registered User",
                "goal": "Apply to join a community and follow the manager approval branch shown in the PDF.",
                "trigger": "User selects Apply to join.",
                "preconditions": "User is viewing community information.",
                "postconditions": "User is joined, or notification rejected/reason path is reached.",
                "main": [
                    "User selects Apply to join.",
                    "System checks Community Manager Approval decision.",
                    "If approval is yes, user reaches Joined state.",
                    "If approval is no, notification rejected/reason is shown.",
                ],
                "alternate": [
                    "Community Manager Approval decision branches to yes/no.",
                ],
                "exceptions": [],
                "rules": ["Approval criteria and manager review steps are not specified in the PDF."],
                "nodes": ["40", "41", "42", "43", "44"],
            },
            {
                "id": "UC-15",
                "name": "View Joined Community Dashboard",
                "heading": "UC-15 - View Joined Community Dashboard",
                "actor": "Community Member",
                "goal": "Open and view a joined community page/dashboard.",
                "trigger": "User clicks on the community group.",
                "preconditions": "User is joined to the community.",
                "postconditions": "Community page/dashboard is shown.",
                "main": [
                    "User clicks on the community group.",
                    "System opens Community page/dashboard.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": [],
                "nodes": ["45", "46"],
            },
            {
                "id": "UC-16",
                "name": "View Community Events and Directory",
                "heading": "UC-16 - View Community Events and Directory",
                "actor": "Community Member",
                "goal": "View community events and directory modules shown in the PDF.",
                "trigger": "User opens Community page/dashboard.",
                "preconditions": "User is in a joined community.",
                "postconditions": "Community event and directory modules are visible.",
                "main": [
                    "User views Event calendar for events, meetings, and webinars.",
                    "User views Past or Upcoming events.",
                    "User views Details or photos.",
                    "User views Member directory, community manager, community member, and friends in community.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not specify member-directory actions beyond viewing these modules."],
                "nodes": ["48", "49", "50", "51", "56", "57", "58", "59"],
            },
            {
                "id": "UC-17",
                "name": "Register for Event",
                "heading": "UC-17 - Register for Event",
                "actor": "Community Member",
                "goal": "Register for an event shown in the community calendar flow.",
                "trigger": "User selects Register on an event.",
                "preconditions": "User is viewing an event.",
                "postconditions": "System shows successfully registered notification.",
                "main": [
                    "System shows new event.",
                    "System shows details such as date, time, and topic.",
                    "User selects Register.",
                    "System sends notification successfully registered.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Capacity, attendance, and webinar joining behavior are not specified in the PDF."],
                "nodes": ["52", "53", "54", "55"],
            },
            {
                "id": "UC-18",
                "name": "Leave Community",
                "heading": "UC-18 - Leave Community",
                "actor": "Community Member",
                "goal": "Leave a joined community.",
                "trigger": "User selects Leave community.",
                "preconditions": "User is in a joined community.",
                "postconditions": "User leaves the community.",
                "main": [
                    "User selects Leave community.",
                    "User clicks leave button.",
                    "System shows confirmation message.",
                    "User confirms Leave.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Role transfer or manager-specific leave behavior is not specified in the PDF."],
                "nodes": ["47", "77", "78", "80"],
            },
        ],
    },
    {
        "title": "Functional Area 4 - Connections and Chats",
        "use_cases": [
            {
                "id": "UC-19",
                "name": "Explore Users",
                "heading": "UC-19 - Explore Users",
                "actor": "Registered User",
                "goal": "Search, browse, sort, filter, and recommend users using options shown in the PDF.",
                "trigger": "User opens Explore.",
                "preconditions": "User is logged in.",
                "postconditions": "User list/profile results are displayed.",
                "main": [
                    "User opens Explore.",
                    "User may Search by keywords, condition, treatment, symptoms, interests, username, or place.",
                    "User may Browse.",
                    "User may Sort by newest, nearest, recently active, similar to me, or most contributions.",
                    "User may Filter by sex, age, gender, location, education, work, friends, symptom, or moderation team.",
                    "System may Recommend based on interests, condition, treatment, symptom, mutual friends, or people user may know.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not specify empty-result behavior."],
                "nodes": ["82", "83", "84", "85", "86", "87"],
            },
            {
                "id": "UC-20",
                "name": "Send Connection Request",
                "heading": "UC-20 - Send Connection Request",
                "actor": "Registered User",
                "goal": "Apply to connect with another user.",
                "trigger": "User selects Apply to connect.",
                "preconditions": "User is viewing a profile.",
                "postconditions": "Request submission notification is created.",
                "main": [
                    "User selects Apply to connect.",
                    "System sends notification for request submission.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Duplicate request handling is not specified in the PDF."],
                "nodes": ["81", "88"],
            },
            {
                "id": "UC-21",
                "name": "Respond to Connection Request",
                "heading": "UC-21 - Respond to Connection Request",
                "actor": "Registered User",
                "goal": "Accept or decline a connection request using the branch shown in the PDF.",
                "trigger": "Connection accepted decision is reached.",
                "preconditions": "A connection request exists.",
                "postconditions": "Accepted or declined notification is shown.",
                "main": [
                    "System checks Connection accepted decision.",
                    "If yes, system sends accepted notification.",
                    "If no, system sends declined notification.",
                ],
                "alternate": [
                    "Connection accepted decision branches to yes/no.",
                ],
                "exceptions": [],
                "rules": ["Detailed requester/receiver screens are not specified in the PDF."],
                "nodes": ["89", "90", "92"],
            },
            {
                "id": "UC-22",
                "name": "Manage Friends Lists",
                "heading": "UC-22 - Manage Friends Lists",
                "actor": "Registered User",
                "goal": "View and use friend list categories shown in the PDF.",
                "trigger": "User opens Friends list.",
                "preconditions": "User has access to connections/friends area.",
                "postconditions": "Friend list categories are visible.",
                "main": [
                    "User opens Friends list.",
                    "System displays Customized List, Restricted List, Acquaintance List, and Close friends List.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF lists categories but does not specify detailed list-management actions."],
                "nodes": ["91", "93"],
            },
            {
                "id": "UC-23",
                "name": "Use Private Chats",
                "heading": "UC-23 - Use Private Chats",
                "actor": "Registered User",
                "goal": "Use direct message or group chat options shown in the PDF.",
                "trigger": "User opens Private Chats.",
                "preconditions": "User is in community/connections context.",
                "postconditions": "User selects chat type.",
                "main": [
                    "User opens Private Chats.",
                    "User selects Direct Message.",
                    "User may select Group Chat.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF names chat types but does not specify message send steps."],
                "nodes": ["61", "62"],
            },
        ],
    },
    {
        "title": "Functional Area 5 - Discussion Forums and Posts",
        "use_cases": [
            {
                "id": "UC-24",
                "name": "Create or Edit Post",
                "heading": "UC-24 - Create or Edit Post",
                "actor": "Community Member",
                "goal": "Create, write, edit, and post discussion content using the visible discussion forum flow.",
                "trigger": "User opens Discussion Forums and selects write/edit or post discussion action.",
                "preconditions": "User is in community engagement/discussion forum context.",
                "postconditions": "Post is successfully posted or remains in the posting flow.",
                "main": [
                    "User opens Discussion Forums.",
                    "User selects Post discussions or Write/edit posts.",
                    "User may post anonymously.",
                    "User selects privacy options such as anyone or community only.",
                    "User may upload photo, videos, documents, or files.",
                    "System shows successfully posted notification.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not specify draft save behavior or exact post form fields."],
                "nodes": ["63", "64", "65", "66", "67", "68", "69", "71"],
            },
            {
                "id": "UC-25",
                "name": "View Comments and Post Interactions",
                "heading": "UC-25 - View Comments and Post Interactions",
                "actor": "Community Member",
                "goal": "View comments and visible interaction information for posts.",
                "trigger": "User opens post comments or interaction details.",
                "preconditions": "Post exists.",
                "postconditions": "Comment and interaction information is shown.",
                "main": [
                    "User views comments.",
                    "User views who likes, saves, or shares the posts.",
                    "System shows people likes, saved, comments/replies the posts.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not specify privacy limits for interaction visibility."],
                "nodes": ["70", "71"],
            },
            {
                "id": "UC-26",
                "name": "Engage with Post",
                "heading": "UC-26 - Engage with Post",
                "actor": "Community Member",
                "goal": "Use post engagement actions shown in the PDF.",
                "trigger": "User interacts with a post or reply.",
                "preconditions": "User is viewing a post.",
                "postconditions": "Selected post engagement action occurs.",
                "main": [
                    "User replies to or comments on a post with text, photo, or video.",
                    "User may tag user on replies.",
                    "User may like posts.",
                    "User may save/follow posts.",
                    "User may share posts.",
                    "User may mark post helpful.",
                    "User may view who likes replies.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF lists these actions but does not specify permission or notification rules for each action."],
                "nodes": ["73", "76"],
            },
            {
                "id": "UC-27",
                "name": "Resolve Question with Best Answer",
                "heading": "UC-27 - Resolve Question with Best Answer",
                "actor": "Community Member",
                "goal": "Mark a question as solved and choose a best answer.",
                "trigger": "User uses question solved or best answer controls.",
                "preconditions": "Question/reply context exists.",
                "postconditions": "Question solved or best answer action is applied.",
                "main": [
                    "User selects Question solved.",
                    "User chooses a best answer for the post.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not specify who is allowed to choose the best answer."],
                "nodes": ["72", "79"],
            },
            {
                "id": "UC-28",
                "name": "Cancel or Delete Posting",
                "heading": "UC-28 - Cancel or Delete Posting",
                "actor": "Community Member",
                "goal": "Cancel or delete a posting using the visible warning/reason flow.",
                "trigger": "User selects cancel/delete posting.",
                "preconditions": "User is in posting flow.",
                "postconditions": "System shows posts will be deleted warning or reasoning.",
                "main": [
                    "User selects Cancel/delete posting.",
                    "System shows posts will be deleted warning or reasoning.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not clarify whether this applies to drafts, published posts, or both."],
                "nodes": ["74"],
            },
            {
                "id": "UC-29",
                "name": "Note Inappropriate or Offensive Content",
                "heading": "UC-29 - Note Inappropriate or Offensive Content",
                "actor": "Community Member",
                "goal": "Use the inappropriate/offensive content note shown in the PDF.",
                "trigger": "User selects note for inappropriate or offensive content.",
                "preconditions": "User is viewing discussion content.",
                "postconditions": "Moderation-related path is referenced.",
                "main": [
                    "User creates note for inappropriate or offensive content.",
                    "PDF references Moderation team in the broader flow.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["Moderator review actions are not specified in the PDF."],
                "nodes": ["72", "87"],
            },
            {
                "id": "UC-30",
                "name": "Explore Discussions",
                "heading": "UC-30 - Explore Discussions",
                "actor": "Community Member",
                "goal": "Explore and search discussion posts.",
                "trigger": "User opens Discussion list or Explore discussion.",
                "preconditions": "User has access to discussion list.",
                "postconditions": "Discussion results are shown.",
                "main": [
                    "User opens Discussion list.",
                    "User opens Explore discussion.",
                    "User may Search by keywords, condition, treatment, or symptom.",
                    "System may use user information such as age, gender, or location.",
                    "User may Sort by latest/newest/recent, recently active, most relevant, popular/most liked, unanswered, or most saved.",
                    "User may Filter by user saved or not, date range, posted by friends/connections, or category.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": ["The PDF does not specify empty-result behavior."],
                "nodes": ["94", "95", "96", "97", "98"],
            },
        ],
    },
    {
        "title": "Functional Area 6 - Notifications and Engagement Settings",
        "use_cases": [
            {
                "id": "UC-31",
                "name": "View and Manage Notifications",
                "heading": "UC-31 - View and Manage Notifications",
                "actor": "Registered User",
                "goal": "View notification outputs and manage notification settings shown in the PDF.",
                "trigger": "A notification is produced or user opens Notification settings.",
                "preconditions": "User is logged in or involved in a flow that produces a notification.",
                "postconditions": "Notification is shown, or a notification setting category is selected.",
                "main": [
                    "System shows notifications for request submission, accepted, declined, rejected/reason, successfully registered, or post activity.",
                    "User opens Engagement.",
                    "User opens Notification settings.",
                    "User sees Frequency, Subscription, Posting, Saved/following posts, Connection, Community, and System options.",
                    "User selects a notification category.",
                ],
                "alternate": [],
                "exceptions": [],
                "rules": [
                    "Read/unread behavior is not specified in the PDF.",
                    "Detailed save behavior for notification settings is not specified in the PDF.",
                ],
                "nodes": ["42", "44", "55", "75", "88", "90", "92", "99", "100", "101"],
            },
        ],
    },
]


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    for name in ["Normal", "Heading 1", "Heading 2", "Title"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title_run = title.add_run("Process Flow - Use Cases")
    title_run.bold = True
    title_run.font.size = Pt(20)
    add_para(doc, "Source: Toni EDIT Process Flow 2.pdf")
    add_para(doc, "Scope: Only functions visible in the PDF are documented. If behavior is not visible in the PDF, it is marked as not specified.")

    doc.add_paragraph("Sample Use Case Format", style="Heading 1")
    add_label(doc, "Use Case ID", "UC-XX")
    add_label(doc, "Use Case Name", "Name from PDF-visible function")
    add_label(doc, "Primary Actor", "Actor visible or implied by the PDF label")
    add_label(doc, "Goal", "What the user/system is trying to do")
    add_label(doc, "Trigger", "Visible starting action or node")
    add_label(doc, "Preconditions", "Only what is visible or necessary from prior nodes")
    add_label(doc, "Postconditions", "Visible end state or next node")
    add_section_label(doc, "🔄 Main Flow:")
    add_numbered(doc, ["Step one", "Step two"])
    add_section_label(doc, "⚠️ Alternate Flows:")
    add_bullets(doc, ["A1: Visible decision path, if shown"])
    add_section_label(doc, "❌ Exceptions:")
    add_bullets(doc, [])
    add_section_label(doc, "📋 Business Rules / Notes:")
    add_bullets(doc, ["Use 'Not specified in PDF' when details are not shown."])
    add_label(doc, "🔗 Related Screens / Index IDs", "01, 02...")

    for area in areas:
        doc.add_paragraph(area["title"], style="Heading 1")
        for uc in area["use_cases"]:
            add_use_case(doc, uc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
